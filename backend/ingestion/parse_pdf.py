# Parse PDF 
from PyPDF2 import PdfReader
from docx import Document
import os
import sys
from pathlib import Path

# Add the backend directory to the Python path for progress service
backend_dir = Path(__file__).parent.parent
if str(backend_dir) not in sys.path:
    sys.path.insert(0, str(backend_dir))

try:
    from app.services.progress_service import progress_service
except ImportError:
    # Fallback if progress service is not available
    progress_service = None

def parse_document(path, file_id=1, progress_file_id=None):
    """
    Parse either PDF or DOCX files based on file extension
    """
    file_extension = os.path.splitext(path)[1].lower()
    chunks = []
    
    if file_extension == '.pdf':
        chunks = parse_pdf(path, file_id, progress_file_id)
    elif file_extension == '.docx':
        chunks = parse_docx(path, file_id, progress_file_id)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}. Only PDF and DOCX files are supported.")
    
    return chunks

def parse_pdf(path, file_id=1, progress_file_id=None):
    """Parse PDF files using PyPDF2"""
    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, "Opening PDF file...")
    
    reader = PdfReader(path)
    total_pages = len(reader.pages)
    chunks = []
    
    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, f"Parsing {total_pages} pages...", 10)
    
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        chunks.append({"file_id": file_id, "page": i, "text": text})
        
        # Update progress for each page
        if progress_service and progress_file_id:
            page_progress = ((i + 1) / total_pages) * 90  # 90% of step 1
            progress_service.update_progress(
                progress_file_id, 
                1, 
                f"Parsed page {i + 1} of {total_pages}", 
                10 + page_progress
            )
    
    return chunks

def parse_docx(path, file_id=1, progress_file_id=None):
    """Parse DOCX files using python-docx"""
    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, "Opening DOCX file...")
    
    doc = Document(path)
    chunks = []
    
    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, "Extracting text from document...", 20)
    
    # Extract text from all paragraphs
    full_text = []
    total_paragraphs = len(doc.paragraphs)
    
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():  # Only add non-empty paragraphs
            full_text.append(paragraph.text)
        
        # Update progress for paragraphs
        if progress_service and progress_file_id and total_paragraphs > 0:
            para_progress = ((i + 1) / total_paragraphs) * 70  # 70% of step 1
            progress_service.update_progress(
                progress_file_id, 
                1, 
                f"Processing paragraph {i + 1} of {total_paragraphs}", 
                20 + para_progress
            )
    
    # Combine all text and create a single chunk (or you can split by paragraphs)
    combined_text = '\n'.join(full_text)
    chunks.append({"file_id": file_id, "page": 0, "text": combined_text})
    
    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, "DOCX parsing completed", 100)
    
    return chunks
